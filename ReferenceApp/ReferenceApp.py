#%% [markdown]
# In this cell, I would include some useful resources that was referred to build this project.
# [python-docx] make footer text smaller: https://python-docx.readthedocs.io/en/latest/user/hdrftr.html
# Consult a few similar project for example: 
# https://stackoverflow.com/questions/69428700/how-to-scrape-full-paper-citation-from-google-scholar-search-results-python
# https://medium.com/@nandinisaini021 scraping-publications-of-aerial-image-research-papers-on-google-scholar-using-python-a0dee9744728
# https://stackoverflow.com/questions/62414552/scraping-citation-text-from-pubmed-search-results-with-beautifulsoup-and-python
# https://nexus.od.nih.gov/all/2015/08/31/pmid-vs-pmcid-whats-the-difference/
# Understanding stdin, stdout, stderr: https://stackoverflow.com/questions/3385201/confused-about-stdin-stdout-and-stderr



#%%
# Import Relevant libraries
from textwrap import indent
import pandas as pd # this library organises data frame 
import requests 
from bs4 import BeautifulSoup as bs
import lxml
import json
import docx # this package get text form word document
from docx import Document
from docx.shared import RGBColor
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import wx

#%%
pd.set_option('display.max_row', None)
pd.set_option('display.max_column', None)
pd.set_option('display.width', 1000)

# %%
# read in excel list of publications, store as journal data frame and abstract data frame
journal_df = pd.read_excel('/Users/yi-chunwang/OneDrive - Perspectum Ltd/Work_Repo/ReferenceApp/Perspectum Publications Tracker.xlsx', sheet_name = 0, header=0)



# %% 
abstract_df = pd.read_excel('/Users/yi-chunwang/OneDrive - Perspectum Ltd/Work_Repo/ReferenceApp/Perspectum Publications Tracker.xlsx', sheet_name=1, header=0)




#%%
# Add the summary paragrph into publication tracker.
word_master = docx.Document('/Users/yi-chunwang/OneDrive - Perspectum Ltd/Work_Repo/ReferenceApp/Perspectum Publications List Master - LATEST.docx')
all_paragraphs = word_master.paragraphs

try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML
import zipfile
import io
import requests    

def get_docx_text(path):
    """Take the path of a docx file as argument, return the text in unicode."""

    WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    PARA = WORD_NAMESPACE + 'p'
    TEXT = WORD_NAMESPACE + 't'

    document = zipfile.ZipFile(path)
    xml_content = document.read('word/document.xml')
    document.close()
    tree = XML(xml_content)

    paragraphs = []
    for paragraph in tree.iter(PARA):
        texts = [n.text for n in paragraph.iter(TEXT) if n.text]
        if texts:
            paragraphs.append(''.join(texts))

    return '\n\n'.join(paragraphs)

para_list = get_docx_text('/Users/yi-chunwang/OneDrive - Perspectum Ltd/Work_Repo/ReferenceApp/Perspectum Publications List Master - LATEST.docx').split('\n\n')

for para in para_list:
    if para.startswith('http'):
        index = para_list.index(para)
        para_list[index-1] = para_list[index-1]+para_list[index]
        para_list.remove(para_list[index-1])
# print(para_list)

summary_dict = {'ref':[], 'summary':[]}

for para in para_list:
    if re.search(r'(2\d{3})', para) != None:
        summary_dict['ref'].append(para)
    else:
        if para == " ":
            continue
        else:
            summary_dict['summary'].append(para)
summary_df = pd.DataFrame(summary_dict)




#%%
# create anchor to merge summary dataframe with journal dataframe

Starttitle = []
for item in summary_df['ref']:
    index = item.find(').')
    slice = item[index+2: index+2+60].lstrip() # slice out a section that is more than 50 characters and cut the white space at the begining
    Starttitle.append(slice[0:50].lower().replace('-', ' ').replace('/', ' ').replace(',', ' ').replace('–', ' ')) # store the first 50 character into the Starttitle column
summary_df['Starttitle'] = Starttitle

Starttitle2 = []
for item in journal_df['Title']:
    if type(item) is float:
        Starttitle2.append(item)
    else:
        Starttitle2.append(item.lstrip()[0:50].lower().replace('-', ' ').replace('/', ' ').replace(',', ' ').replace('–', ' '))
journal_df['Starttitle'] = Starttitle2

# merge two dataframes based on the anchor "Starttitle"
new_journal_df = journal_df.merge(summary_df, how='outer')
# new_journal_df.to_excel('Published and Accepted Studies with Summary.xlsx', sheet_name='To fix match')


 
# %%
# Scrap out the citation detail 
headers  = {'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.114 Safari/537.36'}

# get citation function
def get_citation(query):
    url = 'https://pubmed.ncbi.nlm.nih.gov/?term='+query
    response = requests.get(url, headers=headers)
    soup = bs(response.text, features='lxml') # this return the lxml page of the initial search (search from title)
    redirect = soup.find('div', class_='article-page')
    if redirect != None:
        # For re-direct into article
        pmid = redirect['data-article-pmid']
        second_url = 'https://pubmed.ncbi.nlm.nih.gov/'+pmid+'/citations/'
        second_response = requests.get(second_url, headers=headers) # request to get in the paper's Pubmed page
        soup2 = second_response.json() # render the response in json format
        apa_orig = soup2['apa']['orig'] # slice out the apa formatted 
        try:
            start_index = apa_orig.index('.,')  # some studies have single author
            end_index = apa_orig.index('(')
            modify = apa_orig.replace(apa_orig[start_index+3:end_index], 'et al. ') 
            return modify
        except:
            modify = apa_orig
            return modify
    else:
        first_article = soup.find('article', attrs={"class":"full-docsum", "data-rel-pos":"1"})
        if first_article != None:
            # For study that match on pubmed and with explanation box available and
            # for study that match but no explanation box available 
            # get the 1st position article and find the 1st position article
            pmid = first_article.find('span', class_='docsum-pmid').get_text()
            #print(pmid)
            # the issue with this logic is that we assume the 1st article is the best match
            second_url = 'https://pubmed.ncbi.nlm.nih.gov/'+pmid+'/citations/'
            second_response = requests.get(second_url, headers=headers) # request to get in the paper's Pubmed page
            soup2 = second_response.json() # render the response in json format
            apa_orig = soup2['apa']['orig'] # slice out the apa formatted 
            try:
                start_index = apa_orig.index('.,')  # some studies have single author
                end_index = apa_orig.index('(')
                modify = apa_orig.replace(apa_orig[start_index+3:end_index], 'et al. ') 
                return modify
            except:
                modify = apa_orig
                return modify
        else:
            # For study with no matach at all, try search with DOI
            raise Exception('not searchable. use DOI') # use exception message to allow trying the get_citation_doi() function
            exit()



#%%
# getting citation with DOI function
def get_citation_doi(search_doi):
    url = 'https://pubmed.ncbi.nlm.nih.gov/?term='+search_doi
    response = requests.get(url, headers=headers)
    soup = bs(response.text, features='lxml')
    pmid = soup.find('span', class_ = "docsum-pmid").get_text(strip = True)
    second_url = 'https://pubmed.ncbi.nlm.nih.gov/'+pmid+'/citations/' 
    second_response = requests.get(second_url, headers=headers)
    soup2 = second_response.json()
    apa_orig = soup2['apa']['orig']
    start_index = apa_orig.index('.,')
    end_index = apa_orig.index('(')
    modify = apa_orig.replace(apa_orig[start_index+3:end_index], 'et al. ') 
    return modify




#%%
# Tested function utility
# Test on 5 dataset, represent 5 scenarios
# soup: match, 1 article found. [Multiparametric magnetic resonance for the non-invasive diagnosis of liver disease.]
# soup2: match, directly into article page.[ Sex-specific differences in hepatic fat oxidation and synthesis may explain the higher propensity for NAFLD in men.]
# soup3: match, 2 articel found. [ Characterisation of liver fat in the UK Biobank cohort.]
# soup4: match, no explanation box. [Multiparametric magnetic resonance imaging for early detection of diffuse liver disease.]
# soup5: no match at all. [ The Effect of Multiparametric Magnetic Resonance Imaging in Standard of Care for Non-alcoholic Fatty Liver Disease: Protocol for a Randomised Control Trial.]

for i in range(len(new_journal_df['Title'])):
    try: 
        get_citation(new_journal_df['Title'][i].strip())
        print(i, 'study found with get_citation function')
    except: 
        search_doi = new_journal_df.DOI[i].replace('https://doi.org/', '')
        get_citation_doi(search_doi)
        print(i, 'study found with get_citation_doi function')

redirect = soup.find('div', class_='article-page')
if redirect != None:
    # For soup2, direct into article
    pmid = redirect['data-article-pmid']
    print(pmid)
else:
    first_article = soup.find('article', attrs={"class":"full-docsum", "data-rel-pos":"1"})
    if first_article != None:
        # For soup and soup3, match with explanation box available.
        # For soup4, match but no explanation box, get the 1st position article.
        # all of these scenarios, find the 1st position article
        pmid = first_article.find('span', class_='docsum-pmid').get_text()
        print(pmid)
        # the issue with this logic is that we assume the 1st article is the best match
    else:
        # For soup5, no matach at all
        print('not searchable. use DOI')




#%%
# defiine styles functions
def study_style(modify):
    run = document.add_paragraph(style = 'List Paragraph').add_run(modify)
    font = run.font
    font.bold = True
    font.name = 'Proxima Nova'


def summary_style(modify):
    paragraph = document.add_paragraph(style = 'Normal (Web)')
    run = paragraph.add_run(modify)
    font = run.font
    font.name = 'Proxima Nova'

#%% Write into Doc before formating
#Sorting descending from dataframe did not work -> sorting to be controled by input file


document = Document(docx = '/Users/yi-chunwang/Work_Repo/ReferenceApp/Publication Template.docx')
style = document.styles 

section = document.sections[0]
section.left_margin, section.right_margin, section.top_margin, section.bottom_margin = Cm(1.77), Cm(1.77), Cm(1.77), Cm(1.77) # set margin to narrow

# # set header with logo -> this will be controled with template
# header = section.header
# paragraph = header.paragraphs[0]
# paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
# logo_run = paragraph.add_run()
# # logo_path = input('Please choose a Logo path:')
# # logo_run.add_picture(logo_path, width=Cm(6))
# logo_run.add_picture('/Users/yi-chunwang/OneDrive - Perspectum Ltd/Work_Repo/ReferenceApp/Logo Perspectum_RGB_NoTM.png', width=Cm(6)) # 

# allow user to put in document title -> this will also be controled with template
# heading = input('Type in the document title:')
# document.add_heading(heading, 0)


# Writing reference and summary into the document
for i in range(len(new_journal_df.Title)):
    try:
        query = new_journal_df.Title[i].rstrip().lstrip().replace(' ','+') # this prepare the title to be searchable
        try:
            # get citation with get_citation function
            modify = get_citation(query) # define the variable modify again from the returned value of get_citation function
            study_style(modify)
            if pd.notna(new_journal_df['summary'][i]) is True:
                summary_style(new_journal_df['summary'][i])
            else:
                print('no summary for study:', new_journal_df.Title[i])
            print(i, 'written')
        except:
            # get citation with get_citation_doi function
            search_doi = new_journal_df.DOI[i].replace('https://doi.org/', '')
            modify = get_citation_doi(search_doi)
            study_style(modify)
            if pd.notna(new_journal_df['summary'][i]) is True:
                summary_style(new_journal_df['summary'][i])
            else:
                print('no summary for study:', new_journal_df.Title[i])
            print(i, 'written')
    except:
        print('Can not get citation for the study in index:', i)
        if pd.isna(new_journal_df.Title[i]) is True:
            print('There is no title for the study')
        else:
            if (new_journal_df.Status[i].strip() == 'Published') or (new_journal_df.Status[i] == 'Accepted'):
                print('Published but not match with pubmed') # published, not match on pubmed
            else:
                print('Please check the status of the study:', new_journal_df.Status[i]) # not published
        
document.save('demo1.docx')        



#%%
# Test wriiting into Doc

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH # import the python-docx property setting WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

document = docx.Document(docx = '/Users/yi-chunwang/Work_Repo/ReferenceApp/Publication Template.docx') # use template
style = document.styles 
# lstyle = styles.add_style('List Number', WD_STYLE_TYPE.PARAGRAPH, builtin=True)
"""
To access the style in the template:
from docx.enum.style import WD_STYLE_TYPE
document = docx.Document(docx = '/Users/yi-chunwang/Work_Repo/ReferenceApp/Publication Template.docx') # use template
styles = document.styles
paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
for style in paragraph_styles:
    print(style.name)
"""
# section = document.sections[0]
# section.left_margin, section.right_margin, section.top_margin, section.bottom_margin = Cm(1.77), Cm(1.77), Cm(1.77), Cm(1.77) # set margin to narrow
# header = section.header
# paragraph = header.paragraphs[0]
# paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
# logo_run = paragraph.add_run()
# logo = logo_run.add_picture('/Users/yi-chunwang/OneDrive - Perspectum Ltd/Work_Repo/ReferenceApp/Logo Perspectum_RGB_NoTM.png', width=Cm(6)) 

# footer = section.footer
# paragraph = footer.paragraphs[0]
# today = datetime.date.today().strftime('%d/%m/%Y')
# contact_info = 'info@perspectum.com | www.perspectum.com | U.S.: (+1) 857 321 8675 | U.K.: (+44) 1865 655343'
# page_num = 'page X of Y'  # paragraph.add_run(style[WD_STYLE.PAGE_NUMBER])
# foot_run = paragraph.add_run(today +'\t'+ contact_info + '\t' + page_num) 

# ref_run = document.add_paragraph(style = 'List Number').add_run(modify)
ref_run = document.add_paragraph(style = 'List Paragraph').add_run('reference of the study blabla ')
font = ref_run.font
font.bold = True
font.name = 'Proxima Nova'
font.color.rgb = RGBColor(0x01, 0x42, 0x7E) 

paragraph = document.add_paragraph(style = 'Normal (Web)')
# paragraph_format = paragraph.paragraph_format
# paragraph_format.left_indent = Cm(0.63)
sum_run = paragraph.add_run('summary pojedopvn')
font = sum_run.font
font.name = 'Proxima Nova'


# duplicate
ref_run = document.add_paragraph(style = 'List Paragraph').add_run('reference of the study2')
font = ref_run.font
font.bold = True
font.name = 'Proxima Nova'
font.color.rgb = RGBColor(0x01, 0x42, 0x7E) 

paragraph = document.add_paragraph(style = 'Normal (Web)')
sum_run = paragraph.add_run('summary2')
font = sum_run.font
font.name = 'Proxima Nova'


document.save('test.docx')

#%%
class Abstract(Publication):
    copyright_year = 2021
    def __init__(self, status, technology, disease, summary, conference): # adding a conference attrobute
        super().__init__(status, technology, disease, summary) # this way the subclass handle the common attributes from the parent class, which is calss Publication
        self.conference = conference

abs_1 = Abstract('Accepted', 'CoverScan', 'long Covid','summary5', 'RSNA')
abs_2 = Abstract('In prep', 'MRCP', 'PBS','summary6', 'ISMRM')



abs_1.apply_copyright()
print(abs_1.summary) # summary5Reserved for2021

pub_1.apply_copyright()
print(pub_1.summary) # summary1Reserved for2022

# see the publication class' copyright year still intact



#%%
# Object Oriented Programming

class Publication:
    
    copyright_year = 2022 
    publication_num = 0 

    def __init__(self, status, technology, disease, summary):
        
        self.status = status
        self.technology = technology
        self.disease = disease
        self.summary = summary 
        Publication.publication_num += 1 
    
    def fullinfo(self):
        return '{} {} {} {}'.format(self.status, self.technology, self.disease, self.summary) 
    
    def apply_copyright(self):
        self.summary = self.summary + 'Reserved for' + str(self.copyright_year) 
        return self.summary

    def __repr__(self):
        return 'Study({}, {}, {})'.format(self.status, self.technology, self.disease)
    

    def __str__(self):
        return '{} - {}'.format(self.technology, self.disease)


    @classmethod # convert a regular method to class method
    def set_copyright_year(cls, year): # cls is the convention
        cls.copyright_year = str(year)
    
    @classmethod # provide a method that people can just add a new study by typing in string
    def from_string(cls, pub_str): # use class method as alternative constructor for new study
        status, technology, disease, summary = pub_str.split('_')
        return cls(status, technology, disease, summary) # return a new instance object 
    
    @staticmethod # the method does not operate on the class and instance
    def update_person(initial):
        person = initial
        return 'Last update by:' + person

pub_1 = Publication('Published', 'LMS', 'HCC', 'summary1')
pub_2 = Publication('Accepted', 'MRCP', 'cholangitis', 'summary2')
print(pub_2)
print(pub_1)



#%%
# create a SocialMedia class and inherit from Publication
# Use this to control whether we cover the study or abstract
class SocialMedia(Publication): 

    def __init__(self, status, technology, disease, summary, CoverStudy = None):
        super().__init__(status, technology, disease, summary)
        if CoverStudy is None:
            self.CoverStudy = []
        else:
            self.CoverStudy = CoverStudy
    
    def post_study(self, Study):
        if Study not in self.CoverStudy:
            self.CoverStudy.append(Study)
    
    def remove_study(self, Study):
        if Study in self.CoverStudy:
            self.CoverStudy.remove(Study)
    
    def covered_study(self):
        for Study in self.CoverStudy:
            print(self.fullinfo(), 'Covered-->', Study.fullinfo())
    
    

Sol1 = SocialMedia('Posted', 'LMS', 'NASH', 'sum', [pub_1])
print(Sol1.disease)
Sol1.post_study(pub_2)
Sol1.remove_study(pub_1)
print(Sol1.covered_study())
# # Sol1.covered_study()

#%%

class Publication:

    def __init__(self, status, technology, disease, summary):
        self.status = status
        self.technology = technology
        self.disease = disease

    @property # this way we can access the summary like an attribute and change relevant attributes
    def summary(self):
        return 'The study has the status and technology:{}, {}'.format(self.status, self.technology)
    
    @property # remember we need to declare the method first 
    def fullinfo(self):
        return '{},{},{}'.format(self.status, self.technology, self.disease)
    
    @fullinfo.setter # then turn it into setter 
    def fullinfo(self, info_value):
        status, technology, disease = info_value.split(',')
        self.status = status
        self.technology = technology
        self.disease = disease

    @fullinfo.deleter # and deleter

    def fullinfo(self):
        print('Delete the info!')
        self.status = None
        self.technology = None
        self.disease = None


pub_3 = Publication('Accepted', 'MRCP', 'PSB', 'summary xxx')
pub_3.fullinfo = 'Published, MRCP, cholangitis'

del pub_3.fullinfo # still use it like a attribute, but it is actually deleter function
print(pub_3.status)
print(pub_3.technology)
print(pub_3.disease)

# print(pub_3.fullinfo())
# # print(pub_3.coreinfo())

# # say we want to change summary attribute and let the other attributes also change
# pub_3.status = 'Published'
# print(pub_3.status)
# print(pub_3.summary)




#%%
class RefApp(wx.App):
    pass

class RefFrame(wx.Frame):
    pass

class RefPanel(wx.Panel):
    so_far = 0 # class variable
    def __init__(self, Button, Text, StatusBar, ContinueButton, Input, Output, ErrorM):
        self.Button = Button
        self.Text = Text
        self.StatusBar = tatusBar
        self.ContinueButton = ContinueButton
        self.Input = Input
        self.Output = Output
        self.ErrorM = ErrorM # error message

        RefPanel.so_far += 1 # whenever a new tab is created, the progress workflow add 1 step forward

    def ErrorInfo(self):
        return '{} {} {}'.format(self.Input, self.Output, self.ErrorM)
    
    def __repr__(self):
        return 'Study({}, {})'.format(self.Text, self.Button)
    

    def __str__(self):
        return '{} - {}'.format(self.technology, self.disease)
    
    def __add__(self, other):
        return 'self.Text + self.Button'
        NotImplemented # instead of throwing error, jsut not implement the dunder method because other object might be able to handle the dunder method's operation

    pass

# panel_1 = RefPanel()
# panel_2 = RefPanel()
# panel_3 = RefPanel()
# panel_4 = RefPanel()

class ImportPanel(RefPanel):
    super().__init__(Button, Text, StatusBar, ContinueButton, Input, Output, ErrorM, drag_file = None)
    self.drag_file = drag_file
    if drag_file is None:
        self.drag_file = []
    else:
        self.drag_file.append(drag_file)


class Preview_Panel(RefPanel):
    super().__init__(Button, Text, StatusBar, ContinueButton, Input, Output, ErrorM)



class Style_Panel(RefPanel):
    super().__init__(Button, Text, StatusBar, ContinueButton, Input, Output, ErrorM, template = None)
    if template is None:
        print('please select a template word document(.docx)')
    else:
        self.template = template


class GeneratePanel(RefPanel):
    super().__init__(Button, Text, StatusBar, ContinueButton, Input, Output, ErrorM)





#%%
#TODO: build app functions

"""
Input:
import csv from publication tracker
Select Technology (multiple choice)
Select Sttatus (multiple choice)

Select:
Preview a DataFrame to the next tab
drop rows from data frame (optional)
Preview final dataframe for the studies to be write into word file

Style:
Input Title
select Logo for the document
Select color for reference paragraph

Output:
Gwnerate

"""


#%%
# Create functions to filter columms
import inquirer

status_question = [
    inquirer.List('Status',
                message='What is the status of the article?',
                choices=journal_df['Status'].unique().tolist()
    )
]
# select status filter
selected_status_option = inquirer.prompt(status_question)
journal_df['StatusSelect'] = journal_df['Status'].str.strip().str.lower()
status_list =journal_df['StatusSelect'].unique().tolist()




#%%
# select technology to "LMS"
technology_option = journal_df["Technology"].unique()
selected_technology_option = input('selected from the list')

# Generate filtered dataframe
selected_filter = ( journal_df['Status'] == selected_status_option & journal_df['Technology'] == selected_technology_option )
selected_df = journal_df[selected_filter]
#%%

# %%
